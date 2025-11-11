import os
from datetime import datetime
from typing import Dict, List, Any

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# PDF opcional (si está disponible). Gratis.
# pip install reportlab  (opcional)
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors
    from reportlab.lib.units import cm

    REPORTLAB_OK = True
except Exception:
    REPORTLAB_OK = False


def _safe_first5(identifier: str) -> str:
    """
    Devuelve los 5 primeros caracteres visibles del identificador (incluye '0x' si lo hay).
    """
    if not identifier:
        return "XXXXX"
    s = str(identifier).strip()
    return s[:5]


def _risk_color(band: str):
    band = (band or "").upper()
    if band == "ALTO":
        return (220, 53, 69)  # rojo
    if band == "MEDIO":
        return (255, 193, 7)  # ámbar
    return (25, 135, 84)  # verde


def _docx_set_margins(doc: Document, top=2.0, bottom=2.0, left=2.0, right=2.0):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def _docx_add_header(doc: Document, assets_dir: str):
    """
    Cabecera: logo a la izquierda + membrete. Sin dejar una página en blanco.
    """
    header = doc.sections[0].header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    tab = p.add_run()

    logo_path = os.path.join(assets_dir, "synapxis_logo.png")
    if os.path.exists(logo_path):
        try:
            tab.add_picture(logo_path, width=Inches(1.1))
        except Exception:
            pass

    # Membrete
    run = p.add_run("   SYNAPXIS — Departamento de Análisis Blockchain")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10)


def _docx_title(doc: Document, title: str, risk_band: str):
    # Título
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    r = p.add_run(title)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(16)
    r.bold = True

    # Banda de color de riesgo (línea fina)
    r_color = _risk_color(risk_band)
    rule = doc.add_paragraph()
    rule_run = rule.add_run(" ")
    # Usamos subrayado para simular una franja muy fina de color.
    rule_run.font.underline = True
    try:
        rule_run.font.color.rgb = bytes(r_color)
    except Exception:
        pass


def _docx_kv(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_after = Pt(1)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r1.font.size = Pt(10)
    r2 = p.add_run(value)
    r2.font.name = "Arial"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r2.font.size = Pt(10)


def _docx_section_title(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(12)
    r.bold = True


def _docx_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style=None)
    p.style = doc.styles["List Bullet"]
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(10)


def _ensure_outputs_dir():
    outdir = os.path.join(os.getcwd(), "outputs")
    if not os.path.isdir(outdir):
        os.makedirs(outdir, exist_ok=True)
    return outdir


def generate_docx_and_maybe_pdf(identifier: str,
                                meta: Dict[str, Any],
                                events: List[Dict[str, Any]],
                                transfers: List[Dict[str, Any]]):
    """
    Genera DOCX (si hay reportlab también PDF). Fondo blanco, cabecera con logo,
    márgenes correctos, y aclaraciones N/D y N/A.
    """
    outdir = _ensure_outputs_dir()
    assets_dir = os.path.join(os.getcwd(), "assets")

    # Nombre corto con 5 primeras cifras
    short = _safe_first5(identifier)
    base_name = f"Informe_Wallet_Synapxis_{short}_FINAL"
    docx_path = os.path.join(outdir, f"{base_name}.docx")
    pdf_path = os.path.join(outdir, f"{base_name}.pdf")

    # ===== DOCX =====
    doc = Document()
    _docx_set_margins(doc, 2, 2, 2, 2)
    _docx_add_header(doc, assets_dir)

    # PORTADA (sin salto de página)
    _docx_title(doc, "Informe de Análisis — Wallet / Transacción", meta.get("risk_band", "BAJO"))

    # 1. Resumen general
    _docx_section_title(doc, "1. Resumen general")
    _docx_kv(doc, "Identificador", str(identifier))
    _docx_kv(doc, "Red", str(meta.get("network", "N/D")))
    _docx_kv(doc, "Bloque", str(meta.get("block", "N/D")))
    _docx_kv(doc, "Fecha (UTC)", str(meta.get("timestamp", "N/D")))
    _docx_kv(doc, "Estado", str(meta.get("status", "N/D")))
    _docx_kv(doc, "Riesgo", f"{meta.get('risk_score', 'N/D')}/100 ({meta.get('risk_band', 'N/D')})")

    # 2. Datos básicos
    _docx_section_title(doc, "2. Datos básicos")
    _docx_kv(doc, "Origen (from)", str(meta.get("from", "N/D")))
    _docx_kv(doc, "Destino (to)", str(meta.get("to", "N/D")))
    if meta.get("from_balance_trx") is not None:
        _docx_kv(doc, "Balance TRX", str(meta.get("from_balance_trx")))
    if meta.get("fee_native") is not None:
        _docx_kv(doc, "Fee (nativo)", str(meta.get("fee_native")))
    if meta.get("fee_usd") is not None:
        _docx_kv(doc, "Fee (USD aprox.)", str(meta.get("fee_usd")))

    # 2.b Balances (si existen en meta)
    balances = (meta or {}).get("balances")
    if balances:
        _docx_section_title(doc, "2.b Balances")
        try:
            # Caso dict simple {token: cantidad}
            if isinstance(balances, dict) and "raw" not in balances:
                for k, v in balances.items():
                    _docx_kv(doc, str(k), str(v))
            else:
                # Estructuras complejas: imprime JSON truncado
                from json import dumps
                txt = dumps(balances, indent=2)
                for line in txt[:4000].splitlines():  # truncado defensivo
                    _docx_bullet(doc, line)
        except Exception:
            _docx_bullet(doc, "No se pudo formatear 'balances' (estructura no estándar).")

    # 2.c Direcciones destino detectadas (si las hay)
    targets = (meta or {}).get("targets") or []
    if isinstance(targets, list) and targets:
        _docx_section_title(doc, "2.c Destinos detectados (recientes)")
        for addr in targets[:50]:  # tope por seguridad
            _docx_bullet(doc, f"{addr}")

    # 3. Interpretación técnica
    _docx_section_title(doc, "3. Interpretación técnica")
    notes = []
    if meta.get("flow_label"):
        notes.append(f"Flujo detectado: {meta.get('flow_label')}.")
    if meta.get("network"):
        notes.append(f"Análisis realizado sobre {meta.get('network').upper()}.")
    if not transfers:
        notes.append("No se detectaron transferencias de tokens en esta captura.")
    if not notes:
        notes.append("Sin observaciones técnicas adicionales en esta captura.")
    for n in notes:
        _docx_bullet(doc, n)

    # 4. Evaluación de riesgo
    _docx_section_title(doc, "4. Evaluación de riesgo")
    reasons = meta.get("risk_reasons") or ["Sin señales relevantes con las reglas actuales."]
    for r in reasons:
        _docx_bullet(doc, r)

    # 5. Conclusión interpretativa
    _docx_section_title(doc, "5. Conclusión interpretativa")
    concl = doc.add_paragraph()
    concl_run = concl.add_run(
        f"Nivel de riesgo global: {meta.get('risk_band', 'N/D')}. "
        f"Puntuación estimada: {meta.get('risk_score', 'N/D')}/100. "
        "Este resultado combina coincidencias en listas públicas (si disponibles) "
        "y heurísticas on-chain (flujo, actividad, uso de contratos/tokens)."
    )
    concl_run.font.name = "Arial"
    concl_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    concl_run.font.size = Pt(10)

    # 6. Recomendaciones inmediatas
    _docx_section_title(doc, "6. Recomendaciones inmediatas")
    recs = [
        "Verificar manualmente el destino si el riesgo es MEDIO/ALTO.",
        "Evitar interacción directa con contratos o direcciones marcadas como SCAM/PONZI.",
        "Si la wallet es operativa propia, habilitar alertas y límites de gasto.",
    ]
    for r in recs:
        _docx_bullet(doc, r)

    # 7. Aclaraciones (N/D y N/A)
    _docx_section_title(doc, "7. Aclaraciones")
    aclar = [
        "N/D: No disponible en la captura o no aplicable al tipo de entrada.",
        "N/A: No aplica para el contexto del dato solicitado.",
        "Este informe refleja la información recabada en el momento de la consulta. "
        "La calificación de riesgo puede variar con nuevos eventos en la red."
    ]
    for a in aclar:
        _docx_bullet(doc, a)

    # Pie corporativo
    footer = doc.sections[0].footer
    pf = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    runf = pf.add_run(
        f"Emitido por: Synapxis — Departamento de Análisis Blockchain • "
        f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC\n"
        "Documento preliminar — uso interno o académico. No constituye prueba pericial."
    )
    runf.font.name = "Arial"
    runf._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    runf.font.size = Pt(8)
                                    
        # Pie corporativo
    footer = doc.sections[0].footer
    pf = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    runf = pf.add_run(
        f"Emitido por: Synapxis — Departamento de Análisis Blockchain • "
        f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC\n"
        "Documento preliminar — uso interno o académico. No constituye prueba pericial."
    )
    runf.font.name = "Arial"
    runf._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    runf.font.size = Pt(8)

       # === BLOQUE FUSIONAI ===
    from app.services.fusion_ai import enrich_with_ai

    # 8. Análisis asistido por IA
    try:
       ai_summary = enrich_with_ai(meta, events, transfers)
    if ai_summary:
        _docx_section_title(doc, "8. Análisis asistido por IA (Synapxis FusionAI)")
        _docx_bullet(doc, ai_summary)
    except Exception as e:
        _docx_section_title(doc, "8. Análisis asistido por IA (Synapxis FusionAI)")
        _docx_bullet(doc, f"[FusionAI error] {e}")

    # === FIN BLOQUE FUSIONAI ===

    # Guardar DOCX
    doc.save(docx_path)

    # ===== PDF (opcional) =====
    pdf_generated = False
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        c = canvas.Canvas(pdf_path, pagesize=A4)
        c.setFont("Helvetica", 10)
        textobject = c.beginText(40, 800)
        textobject.textLine(
            f"Emitido por: Synapxis — Departamento de Análisis Blockchain • "
            f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC"
        )
        textobject.textLine("Documento preliminar — uso interno o académico. No constituye prueba pericial.")
        c.drawText(textobject)
        c.showPage()
        c.save()
        pdf_generated = True
    except Exception:
        pdf_generated = False

    # Archivos de salida
    files = [{"name": os.path.basename(docx_path), "path": docx_path}]
    if pdf_generated:
        files.append({"name": os.path.basename(pdf_path), "path": pdf_path})
    return files
